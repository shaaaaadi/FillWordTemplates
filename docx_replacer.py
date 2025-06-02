from docx import Document
from docx.shared import Inches
import re
import time


class DocxReplacer:
    def __init__(self, input_file, output_file, replacements):
        self.input_file = input_file
        self.output_file = output_file
        self.replacements = replacements
        self.doc = Document(input_file)

    def _apply_text_replacement(self, text, text_replacements):
        for key, val in text_replacements.items():
            if key in text:
                text = text.replace(key, val)
        return text

    from docx.shared import Inches
    def _replace_in_paragraphs(self, text_replacements, image_replacements):
        def remove_placeholder_runs(paragraph, placeholder):
            runs_to_remove = []
            i = 0
            buffer = ""
            start_index = -1

            # Step 1: Reconstruct full text and track run indices
            for run in paragraph.runs:
                buffer += run.text
                if placeholder in buffer and start_index == -1:
                    # Found the start of the placeholder
                    start_index = i
                i += 1

            if start_index == -1:
                return  # Placeholder not found

            # Step 2: Remove only runs that include the placeholder
            i = 0
            matched = False
            remaining_placeholder = placeholder
            for run in paragraph.runs:
                if not matched and remaining_placeholder and remaining_placeholder.startswith(run.text):
                    runs_to_remove.append(run)
                    remaining_placeholder = remaining_placeholder[len(run.text):]
                    if remaining_placeholder == "":
                        matched = True
                i += 1

            for run in runs_to_remove:
                run.text = ""
            return paragraph

        for paragraph in self.doc.paragraphs:
            full_text = "".join(run.text for run in paragraph.runs)
            for key, image_path in image_replacements.items():
                if key in full_text:
                    paragraph = remove_placeholder_runs(paragraph, key)
                    run = paragraph.add_run()
                    run.add_picture(image_path, width=Inches(1))
                    full_text = full_text.replace(key, "")

            replaced_text = self._apply_text_replacement(full_text, text_replacements)
            if replaced_text != full_text:
                for run in paragraph.runs:
                    run.text = ""
                if paragraph.runs:
                    paragraph.runs[0].text = replaced_text
                else:
                    paragraph.add_run(replaced_text)

    def _replace_in_tables(self, text_replacements, image_replacements):
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        full_text = "".join(run.text for run in paragraph.runs)
                        for key, image_path in image_replacements.items():
                            if key in full_text:
                                for run in paragraph.runs:
                                    run.text = ""
                                run = paragraph.add_run()
                                run.add_picture(image_path, width=Inches(2))
                                full_text = full_text.replace(key, "")

                        replaced_text = self._apply_text_replacement(full_text, text_replacements)
                        if replaced_text != full_text:
                            for run in paragraph.runs:
                                run.text = ""
                            if paragraph.runs:
                                paragraph.runs[0].text = replaced_text
                            else:
                                paragraph.add_run(replaced_text)

    def _replace_in_textboxes(self, text_replacements):
        for shape in self.doc.element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent'):
            for text_elem in shape.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                if text_elem.text:
                    text_elem.text = self._apply_text_replacement(text_elem.text, text_replacements)

    def _generate_final_replacements(self):
        text_replacements = {}
        image_replacements = {}
        for item in self.replacements:
            key = item.get("key")
            value = item.get("value")
            type_ = item.get("type")

            if type_ == "string":
                text_replacements[key] = value
            elif type_ == "full_name":
                parts = value.strip().split(maxsplit=1)
                first_name = parts[0]
                surname = parts[1] if len(parts) > 1 else ''
                text_replacements["<first_name>"] = first_name
                text_replacements["<surname>"] = surname
            elif type_ == "id":
                if len(value) == 9:
                    for i, digit in enumerate(value):
                        text_replacements[f"<id_{i+1}>"] = digit
            elif type_ == "signature":
                image_replacements[key] = value

        return text_replacements, image_replacements

    def _clear_unreplaced_placeholders(self):
        placeholder_pattern = r"<[^>]+>"

        for paragraph in self.doc.paragraphs:
            full_text = "".join(run.text for run in paragraph.runs)
            cleaned_text = re.sub(placeholder_pattern, "", full_text)
            if cleaned_text != full_text:
                for run in paragraph.runs:
                    run.text = ""
                if paragraph.runs:
                    paragraph.runs[0].text = cleaned_text
                else:
                    paragraph.add_run(cleaned_text)

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        full_text = "".join(run.text for run in paragraph.runs)
                        cleaned_text = re.sub(placeholder_pattern, "", full_text)
                        if cleaned_text != full_text:
                            for run in paragraph.runs:
                                run.text = ""
                            if paragraph.runs:
                                paragraph.runs[0].text = cleaned_text
                            else:
                                paragraph.add_run(cleaned_text)

        for shape in self.doc.element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent'):
            for text_elem in shape.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                if text_elem.text:
                    text_elem.text = re.sub(placeholder_pattern, "", text_elem.text)

    def run(self):
        text_replacements, image_replacements = self._generate_final_replacements()
        self._replace_in_paragraphs(text_replacements, image_replacements)
        self._replace_in_tables(text_replacements, image_replacements)
        self._replace_in_textboxes(text_replacements)
        self._clear_unreplaced_placeholders()
        self.doc.save(self.output_file)


# Example usage
if __name__ == "__main__":
    input_file = r"C:\word\bekoretKatsen.docx"
    output_file = rf"C:\word\bekoretKatsen_{time.time()}.docx"
    replacements = [
        {"key": "<cusName>", "value": "שאדי", "type": "string"},
        {"key": "<cusPhone>", "value": "0543341222", "type": "string"},
        {"key": "<hnDate>", "value": "14/11/2009", "type": "string"},
        {"key": "<date>", "value": "14/11/2010", "type": "string"},
        {"key": "<kilometer>", "value": "14/11/2009", "type": "string"},
        {"key": "<cn_1>", "value": "798", "type": "string"},
        {"key": "<cn_2>", "value": "30", "type": "string"},
        {"key": "<cn_3>", "value": "102", "type": "string"},
        {"key": "<P1a>", "value": "X", "type": "string"},
        {"key": "<winterSignature>", "value": "c:\\word\\s1.png", "type": "signature"},
        {"key": "<officerSignature>", "value": "c:\\word\\s2.png", "type": "signature"}
    ]

    replacer = DocxReplacer(input_file, output_file, replacements)
    replacer.run()